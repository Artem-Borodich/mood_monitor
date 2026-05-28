# -*- coding: utf-8 -*-
import glob
import os
from docx import Document

root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
paths = [
    p
    for p in glob.glob(os.path.join(root, "*.docx"))
    if ".bak_" not in p and not os.path.basename(p).startswith("~$")
]
doc = Document(paths[0])
for i in range(160, 195):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    t = p.text.strip()
    print(i, repr(p.style.name), t[:100] if t else "(empty)")
