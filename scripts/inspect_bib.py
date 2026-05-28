# -*- coding: utf-8 -*-
import glob
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8")
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def inspect(path: str, label: str) -> None:
    doc = Document(path)
    print(f"\n{'='*60}\n{label}: {os.path.basename(path)}\n{'='*60}")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Список использованных источников":
            start = i
            for j in range(i, min(i + 8, len(doc.paragraphs))):
                pp = doc.paragraphs[j]
                tt = pp.text.strip()
                if not tt and j > i:
                    continue
                pf = pp.paragraph_format
                pr = pp._element.find(qn("w:pPr"))
                ps = pr.find(qn("w:pStyle")) if pr is not None else None
                sid = ps.get(qn("w:val")) if ps is not None else None
                r = pp.runs[0] if pp.runs else None
                print(f"[{j}] {tt[:95]}")
                print(
                    f"    style={pp.style.name} sid={sid} "
                    f"align={pf.alignment} ind={pf.first_line_indent} left={pf.left_indent}"
                )
                print(
                    f"    line={pf.line_spacing_rule} {pf.line_spacing} "
                    f"font={r.font.name if r else None} sz={r.font.size if r else None}"
                )
                if j >= i + 6:
                    break
            break


from diploma_paths import diploma_path, DIPLOMA_FILENAME

ex = [
    p
    for p in glob.glob(os.path.join(ROOT, "*.docx"))
    if os.path.basename(p) != DIPLOMA_FILENAME
    and ".bak" not in p
    and "updated" not in p.lower()
    and not os.path.basename(p).startswith("~$")
]
if not ex:
    raise SystemExit("Пример не найден")
ex = ex[0]

mine = diploma_path()
inspect(ex, "ПРИМЕР")
inspect(mine, "МОЙ")
