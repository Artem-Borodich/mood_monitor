# -*- coding: utf-8 -*-
"""Сравнение оформления диплома и примера."""
from __future__ import annotations

import glob
import os
import sys
from collections import Counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def pt(val) -> str:
    if val is None:
        return "—"
    try:
        return f"{val.pt:.1f} pt"
    except Exception:
        return str(val)


def find_diploma_paths() -> tuple[str, str]:
    from diploma_paths import diploma_path, example_docx_path

    mine = diploma_path()
    ex_path = example_docx_path()
    if not ex_path:
        raise FileNotFoundError("Пример DOCX не найден (diploma/references/)")
    return mine, ex_path


def style_summary(doc: Document) -> dict:
    styles_used: Counter = Counter()
    para_formats: list[dict] = []
    for p in doc.paragraphs:
        name = p.style.name if p.style else "?"
        styles_used[name] += 1
        if p.text.strip() and len(para_formats) < 5:
            pf = p.paragraph_format
            r0 = p.runs[0] if p.runs else None
            para_formats.append(
                {
                    "style": name,
                    "align": str(pf.alignment),
                    "first_indent": pt(pf.first_line_indent),
                    "line_spacing": str(pf.line_spacing),
                    "space_after": pt(pf.space_after),
                    "font": r0.font.name if r0 else None,
                    "size": pt(r0.font.size) if r0 and r0.font.size else None,
                    "text": p.text[:50],
                }
            )
    return {"styles": styles_used, "samples": para_formats}


def heading_styles(doc: Document) -> list[dict]:
    out = []
    for level in range(1, 4):
        try:
            st = doc.styles[f"Heading {level}"]
            f = st.font
            pf = st.paragraph_format
            out.append(
                {
                    "level": level,
                    "font": f.name,
                    "size": pt(f.size),
                    "bold": f.bold,
                    "align": str(pf.alignment),
                    "space_before": pt(pf.space_before),
                    "space_after": pt(pf.space_after),
                }
            )
        except KeyError:
            pass
    return out


def caption_samples(doc: Document) -> list[dict]:
    caps = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Рисунок") or t.startswith("Таблица"):
            r = p.runs[0] if p.runs else None
            caps.append(
                {
                    "text": t[:60],
                    "style": p.style.name,
                    "align": str(p.alignment),
                    "italic": r.font.italic if r else None,
                    "size": pt(r.font.size) if r and r.font.size else None,
                }
            )
    return caps[:5]


def bibliography_info(doc: Document) -> dict:
    in_bib = False
    items = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Список использованных" in t or "СПИСОК ИСПОЛЬЗОВАННЫХ" in t.upper():
            in_bib = True
            continue
        if in_bib:
            if not t:
                continue
            if t.startswith("Приложение") or t.startswith("ПРИЛОЖЕНИЕ"):
                break
            items.append(t[:80])
    refs_in_text = sum(1 for p in doc.paragraphs if "[" in p.text and "]" in p.text)
    return {"count": len(items), "samples": items[:3], "bracket_refs": refs_in_text}


def section_numbering(doc: Document) -> list[str]:
    heads = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style.name.startswith("Heading") and t:
            heads.append(f"{p.style.name}: {t[:50]}")
    return heads[:15]


def analyze(path: str, label: str) -> None:
    doc = Document(path)
    print(f"\n{'='*60}\n{label}: {os.path.basename(path)}\n{'='*60}")
    print("Параграфов:", len(doc.paragraphs))
    ss = style_summary(doc)
    print("\nТоп стилей:")
    for name, cnt in ss["styles"].most_common(12):
        print(f"  {name}: {cnt}")
    print("\nЗаголовки (стиль документа):")
    for h in heading_styles(doc):
        print(f"  H{h['level']}: {h['font']} {h['size']} bold={h['bold']} align={h['align']}")
    print("\nПодписи рисунков/таблиц:")
    for c in caption_samples(doc):
        print(f"  {c}")
    print("\nСписок литературы:", bibliography_info(doc))
    print("\nЗаголовки разделов (первые):")
    for h in section_numbering(doc):
        print(f"  {h}")


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    mine, example = find_diploma_paths()
    analyze(example, "ПРИМЕР")
    analyze(mine, "МОЙ")


if __name__ == "__main__":
    main()
