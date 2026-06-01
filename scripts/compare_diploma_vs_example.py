# -*- coding: utf-8 -*-
"""Сравнение Диплома.docx с примером — отчёт в _compare_report.txt"""
from __future__ import annotations

import glob
import os
import re
import zipfile
from collections import Counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def pt(val) -> str:
    if val is None:
        return "—"
    try:
        return f"{val.pt:.1f} pt"
    except Exception:
        return str(val)


def cm(val) -> str:
    if val is None:
        return "—"
    try:
        return f"{val.cm:.2f} см"
    except Exception:
        return "—"


def has_drawing(p) -> bool:
    return bool(p._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"))


def find_paths() -> tuple[str, str]:
    from diploma_paths import diploma_path, example_docx_path

    mine = diploma_path()
    ex_path = example_docx_path()
    if not ex_path:
        raise FileNotFoundError(
            "Пример не найден (положите .docx-пример в diploma/references/)"
        )
    return mine, ex_path


def section_margins(path: str) -> dict:
    doc = Document(path)
    s = doc.sections[0]
    return {
        "top": cm(s.top_margin),
        "bottom": cm(s.bottom_margin),
        "left": cm(s.left_margin),
        "right": cm(s.right_margin),
        "page_w": cm(s.page_width),
        "page_h": cm(s.page_height),
    }


def footer_count(path: str) -> int:
    with zipfile.ZipFile(path) as z:
        return len([n for n in z.namelist() if "footer" in n])


def style_stats(doc: Document) -> Counter:
    return Counter(p.style.name for p in doc.paragraphs if p.text.strip())


def sample_body_para(doc: Document, style_name: str) -> dict | None:
    for p in doc.paragraphs:
        if p.style.name != style_name:
            continue
        if len(p.text.strip()) < 120:
            continue
        if p.text.strip().startswith(("Рисунок", "Листинг", "1.", "2.")):
            continue
        pf = p.paragraph_format
        r = p.runs[0] if p.runs else None
        return {
            "style": style_name,
            "text": p.text[:55] + "...",
            "align": str(pf.alignment),
            "indent": cm(pf.first_line_indent),
            "line_rule": str(pf.line_spacing_rule),
            "line": str(pf.line_spacing),
            "font": r.font.name if r else None,
            "size": pt(r.font.size) if r and r.font.size else None,
        }
    return None


def heading_sample(doc: Document, level: int) -> dict:
    name = f"Heading {level}"
    for p in doc.paragraphs:
        if p.style.name == name and p.text.strip():
            pf = p.paragraph_format
            r = p.runs[0] if p.runs else None
            return {
                "text": p.text.strip()[:50],
                "align": str(pf.alignment),
                "size": pt(r.font.size) if r and r.font.size else None,
                "bold": r.font.bold if r else None,
                "space_before": pt(pf.space_before),
            }
    return {}


def caption_samples(doc: Document, n: int = 3) -> list[dict]:
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not (t.startswith("Рисунок") or t.startswith("Таблица")):
            continue
        r = p.runs[0] if p.runs else None
        out.append(
            {
                "text": t[:55],
                "style": p.style.name,
                "align": str(p.alignment),
                "italic": r.font.italic if r else None,
                "size": pt(r.font.size) if r and r.font.size else None,
            }
        )
        if len(out) >= n:
            break
    return out


def listing_samples(doc: Document) -> list[dict]:
    out = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Листинг") and "–" in p.text:
            out.append({"title": p.text[:50], "title_style": p.style.name})
            for j in range(1, 3):
                if i + j < len(doc.paragraphs):
                    cp = doc.paragraphs[i + j]
                    cr = cp.runs[0] if cp.runs else None
                    out.append(
                        {
                            "code_style": cp.style.name,
                            "font": cr.font.name if cr else None,
                            "size": pt(cr.font.size) if cr and cr.font.size else None,
                        }
                    )
            return out
    return out


def refs_stats(doc: Document) -> dict:
    all_refs = re.findall(r"\[(\d+)\]", "\n".join(p.text for p in doc.paragraphs))
    nums = [int(x) for x in all_refs]
    return {
        "total": len(all_refs),
        "unique": sorted(set(nums)) if nums else [],
        "max": max(nums) if nums else 0,
    }


def bibliography_block(doc: Document) -> dict:
    bib_idx = None
    items = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "Список использованных" in t:
            bib_idx = i
            heading_style = p.style.name
            continue
        if bib_idx is not None:
            if re.match(r"^\d+\.\s", t):
                pf = p.paragraph_format
                items.append(
                    {
                        "n": int(re.match(r"^(\d+)", t).group(1)),
                        "preview": t[:70],
                        "left": cm(pf.left_indent),
                        "first": cm(pf.first_line_indent),
                    }
                )
            elif t and p.style.name.startswith("Heading") and i > bib_idx + 1:
                break
    nums = [x["n"] for x in items]
    return {
        "heading_idx": bib_idx,
        "heading_style": heading_style if bib_idx else None,
        "count": len(items),
        "order_ok": nums == list(range(1, len(nums) + 1)) if nums else False,
        "first": items[0] if items else None,
        "last": items[-1] if items else None,
        "format_en": bool(items and ": monogr." in items[0]["preview"]),
        "format_ru": bool(items and "–" in items[0]["preview"] and ": monogr." not in items[0]["preview"]),
    }


def structure_issues(doc: Document) -> list[str]:
    issues = []
    headings = [(i, p.text.strip(), p.style.name) for i, p in enumerate(doc.paragraphs) if p.style.name.startswith("Heading") and p.text.strip()]
    for i, (idx, t, st) in enumerate(headings):
        if st == "Heading 2" and t and not re.match(r"^(\d+\.|Введение|Заключение|Список)", t):
            if t == "Характеристика приложения":
                issues.append(f"§ без номера: «{t}» (ожидается 1.1 …)")
    # порядок заключение / список
    concl = bib = None
    for idx, t, st in headings:
        if t == "Заключение":
            concl = idx
        if "Список использованных" in t:
            bib = idx
    if concl is not None and bib is not None and concl < bib:
        issues.append("Заключение идёт ПЕРЕД списком литературы (в примере — после основной части, список в конце)")
    has_toc = any(
        "Оглавление" in p.text or "СОДЕРЖАНИЕ" in p.text.upper()
        for p in doc.paragraphs[:40]
    )
    if not has_toc:
        issues.append("Нет явного раздела «Оглавление» в начале (в примере — оглавление с toc)")
    return issues


def count_figures(doc: Document) -> tuple[int, int]:
    caps = sum(1 for p in doc.paragraphs if p.text.strip().startswith("Рисунок"))
    imgs = sum(1 for p in doc.paragraphs if has_drawing(p))
    return caps, imgs


def main() -> None:
    mine_p, ex_p = find_paths()
    mine = Document(mine_p)
    ex = Document(ex_p)

    lines: list[str] = []
    w = lines.append

    w("=" * 70)
    w("СРАВНЕНИЕ: Диплома.docx vs пример (Нестерович)")
    w("=" * 70)
    w(f"Мой файл: {mine_p}")
    w(f"Пример:   {ex_p}")
    w("")

    w("--- ОБЪЁМ ---")
    w(f"Параграфов: мой {len(mine.paragraphs)}, пример {len(ex.paragraphs)}")
    mc, mi = count_figures(mine)
    ec, ei = count_figures(ex)
    w(f"Подписей «Рисунок»: мой {mc}, пример {ec}")
    w(f"Встроенных изображений: мой {mi}, пример {ei}")
    w("")

    w("--- ПОЛЯ СТРАНИЦЫ (секция 0) ---")
    w(f"Мой:    {section_margins(mine_p)}")
    w(f"Пример: {section_margins(ex_p)}")
    w(f"Колонтитулов (footer*.xml): мой {footer_count(mine_p)}, пример {footer_count(ex_p)}")
    w("")

    w("--- СТИЛИ (топ-8) ---")
    w("Мой: " + ", ".join(f"{k}:{v}" for k, v in style_stats(mine).most_common(8)))
    w("Пример: " + ", ".join(f"{k}:{v}" for k, v in style_stats(ex).most_common(8)))
    w("")

    w("--- ОСНОВНОЙ ТЕКСТ ---")
    mb = sample_body_para(mine, "Normal") or sample_body_para(mine, "List Paragraph")
    eb = sample_body_para(ex, "Документация")
    w(f"Мой:    {mb}")
    w(f"Пример: {eb}")
    w("")

    w("--- ЗАГОЛОВКИ ---")
    for lvl in (1, 2, 3):
        w(f"H{lvl} мой:    {heading_sample(mine, lvl)}")
        w(f"H{lvl} пример: {heading_sample(ex, lvl)}")
    w("")

    w("--- ПОДПИСИ РИСУНКОВ ---")
    w("Мой: " + str(caption_samples(mine)))
    w("Пример: " + str(caption_samples(ex)))
    w("")

    w("--- ЛИСТИНГИ ---")
    w("Мой: " + str(listing_samples(mine)))
    w("Пример: " + str(listing_samples(ex)))
    w("")

    w("--- ССЫЛКИ [N] ---")
    w(f"Мой:    {refs_stats(mine)}")
    w(f"Пример: {refs_stats(ex)}")
    w("")

    w("--- СПИСОК ЛИТЕРАТУРЫ ---")
    w(f"Мой:    {bibliography_block(mine)}")
    w(f"Пример: {bibliography_block(ex)}")
    w("")

    w("--- СТРУКТУРА / СОДЕРЖАНИЕ ---")
    for iss in structure_issues(mine):
        w(f"  ! {iss}")
    w("")

    w("--- НУМЕРАЦИЯ РИСУНКОВ (первые подписи) ---")
    my_caps = [p.text.strip() for p in mine.paragraphs if p.text.strip().startswith("Рисунок")][:8]
    ex_caps = [p.text.strip() for p in ex.paragraphs if p.text.strip().startswith("Рисунок")][:5]
    w("Мой: " + str(my_caps))
    w("Пример: " + str(ex_caps))

    from diploma_paths import workspace_file

    out = workspace_file("reports", "_compare_report.txt")
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(out)


if __name__ == "__main__":
    main()
