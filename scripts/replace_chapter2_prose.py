# -*- coding: utf-8 -*-
"""Replace §2.1–2.3 with prose-only version (no figures, no English filler)."""
from __future__ import annotations

import glob
import os
import shutil
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from chapter2_blocks_prose import CHAPTER2_BLOCKS
from insert_chapter2 import (
    apply_body_format,
    docx_paths,
    insert_blocks_before,
)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def collect_section_paragraphs(doc: Document) -> tuple[list[Paragraph], Paragraph | None]:
    """Paragraphs from 2.1 through line before 2.4."""
    to_remove: list[Paragraph] = []
    anchor: Paragraph | None = None
    in_section = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("2.1 Архитектура"):
            in_section = True
        if in_section:
            if t.startswith("2.4 "):
                anchor = p
                break
            to_remove.append(p)
    return to_remove, anchor


def fix_chapter2_heading(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("2 ") and (
            "Проектирование" in t or "Программная реализация" in t or "API" in t
        ):
            p.text = (
                "2 Проектирование и реализация программной системы "
                "мониторинга настроения (интерфейс «Серенити»)"
            )
            break


def main() -> None:
    paths = docx_paths()
    if not paths:
        raise SystemExit("Файл диплома не найден. Закройте Word.")
    path = [p for p in paths if "_updated" not in p][0]

    backup = path + ".bak_prose_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    body_template = next(
        (p for p in doc.paragraphs if p.style.name == "Normal" and len(p.text.strip()) > 80),
        doc.paragraphs[29],
    )

    to_remove, anchor = collect_section_paragraphs(doc)
    if not anchor or not to_remove:
        raise SystemExit("Раздел 2.1–2.3 не найден в документе.")

    for p in reversed(to_remove):
        remove_paragraph(p)

    insert_blocks_before(anchor, CHAPTER2_BLOCKS, body_template)
    anchor.text = "2.4 Программная реализация интерфейса прикладного программирования"
    fix_chapter2_heading(doc)

    # Remove orphan figure captions if any left in ch.2
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Рисунок 2."):
            remove_paragraph(p)

    try:
        doc.save(path)
        print("Обновлён:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_updated.docx")
        doc.save(alt)
        print("Файл занят Word. Сохранено в:", alt)


if __name__ == "__main__":
    main()
