# -*- coding: utf-8 -*-
"""
Приведение Диплома.docx к оформлению примера (Нестерович Диплом.docx).
Только стили, библиография и ссылки [N] — текст не переписывается.
"""
from __future__ import annotations

import glob
import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from diploma_sources import BIBLIOGRAPHY, SOURCE_KEYWORDS

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

FONT_BODY = "Times New Roman"
SIZE_BODY = Pt(14)
SIZE_CODE = Pt(10)
FIRST_INDENT = Cm(1.25)
HANGING_INDENT = Cm(1.25)


from diploma_paths import diploma_path


def docx_path() -> str:
    return diploma_path()


def has_image(paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def format_body_run(run) -> None:
    run.font.name = FONT_BODY
    run.font.size = SIZE_BODY
    run.font.italic = False
    run.font.bold = False


def apply_body_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = FIRST_INDENT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        format_body_run(run)


def apply_heading_format(paragraph, level: int) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    if level == 1:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        size = Pt(14)
        bold = True
    elif level == 2:
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        size = Pt(14)
        bold = True
    else:
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        size = Pt(14)
        bold = True
    for run in paragraph.runs:
        run.font.name = FONT_BODY
        run.font.size = size
        run.font.bold = bold
        run.font.italic = False


def apply_caption_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = FONT_BODY
        run.font.size = SIZE_BODY
        run.font.italic = False
        run.font.bold = False


def apply_code_format(paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run.font.size = SIZE_CODE
        run.font.italic = False


def apply_listing_title_format(paragraph) -> None:
    apply_body_format(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_bibliography_entry(paragraph, text: str) -> None:
    paragraph.text = text
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0)
    pf.left_indent = HANGING_INDENT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        format_body_run(run)


def is_caption(text: str) -> bool:
    t = text.strip()
    return t.startswith("Рисунок") or t.startswith("Таблица") or t.startswith("Листинг")


def is_listing_title(text: str) -> bool:
    return text.strip().startswith("Листинг") and "–" in text


def heading_level(style_name: str) -> int | None:
    if style_name == "Heading 1":
        return 1
    if style_name == "Heading 2":
        return 2
    if style_name == "Heading 3":
        return 3
    return None


def find_zones(doc: Document) -> tuple[int, int, int]:
    """(start_body, bib_start, end) — индексы параграфов."""
    intro_idx = 0
    bib_idx = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Введение" and p.style.name == "Heading 1":
            intro_idx = i
        if "Список использованных источников" in t:
            bib_idx = i
    return intro_idx, bib_idx, len(doc.paragraphs)


def add_references_to_paragraph(paragraph) -> bool:
    if has_image(paragraph):
        return False
    text = paragraph.text
    if not text.strip() or heading_level(paragraph.style.name):
        return False
    if is_caption(text):
        return False

    existing = {int(x) for x in re.findall(r"\[(\d+)\]", text)}
    to_add: list[int] = []
    lower = text.lower()
    for keywords, num in SOURCE_KEYWORDS:
        if num in existing or num in to_add:
            continue
        if any(kw.lower() in lower for kw in keywords):
            to_add.append(num)

    if not to_add:
        return False

    suffix = "".join(f"[{n}]" for n in sorted(set(to_add)))
    stripped = text.rstrip()
    if stripped.endswith("."):
        new_text = stripped[:-1] + " " + suffix + "."
    elif stripped.endswith(";"):
        new_text = stripped[:-1] + " " + suffix + ";"
    else:
        new_text = stripped + " " + suffix

    paragraph.text = new_text
    apply_body_format(paragraph)
    return True


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def replace_bibliography(doc: Document, bib_start: int) -> None:
    """Заменяет пункты списка литературы после заголовка раздела."""
    item_paras: list = []
    i = bib_start + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            i += 1
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        if t.startswith("Приложение"):
            break
        if re.match(r"^\d+\.\s", t):
            item_paras.append(p)
        i += 1

    for p in item_paras:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    heading = doc.paragraphs[bib_start]
    last = heading
    for idx, entry in enumerate(BIBLIOGRAPHY):
        numbered = f"{idx + 1}. {entry}"
        new_p = insert_paragraph_after(last, numbered)
        apply_bibliography_entry(new_p, numbered)
        last = new_p


def format_document(doc: Document) -> dict[str, int]:
    intro_idx, bib_idx, _ = find_zones(doc)
    stats = {"body": 0, "headings": 0, "captions": 0, "code": 0, "refs": 0}

    for i, p in enumerate(doc.paragraphs):
        if i < intro_idx or i >= bib_idx:
            continue
        if i == bib_idx:
            continue

        t = p.text.strip()
        lvl = heading_level(p.style.name)

        if lvl:
            apply_heading_format(p, lvl)
            stats["headings"] += 1
            continue

        if is_caption(t):
            apply_caption_format(p)
            stats["captions"] += 1
            continue

        if p.style.name == "Normal (Web)" or (
            t and "def " in t or "@router" in t or "import " in t
        ):
            apply_code_format(p)
            stats["code"] += 1
            continue

        if is_listing_title(t):
            apply_listing_title_format(p)
            stats["captions"] += 1
            continue

        if p.style.name in ("Normal", "List Paragraph", "Body Text", "Документация"):
            if add_references_to_paragraph(p):
                stats["refs"] += 1
            apply_body_format(p)
            stats["body"] += 1

    replace_bibliography(doc, bib_idx)
    return stats


def update_normal_style(doc: Document) -> None:
    """Базовый стиль Normal — как «Документация» в примере."""
    try:
        st = doc.styles["Normal"]
    except KeyError:
        return
    st.font.name = FONT_BODY
    st.font.size = SIZE_BODY
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = FIRST_INDENT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5


def main() -> None:
    path = docx_path()
    backup = path + ".bak_format_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    update_normal_style(doc)
    stats = format_document(doc)

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_formatted.docx")
        doc.save(alt)
        print("Файл занят Word. Сохранено:", alt)

    print("Статистика:", stats)
    print("Источников в списке:", len(BIBLIOGRAPHY))


if __name__ == "__main__":
    main()
