# -*- coding: utf-8 -*-
import re

from docx import Document

from diploma_paths import diploma_path, workspace_file

doc = Document(diploma_path())
full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
out = workspace_file("exports", "_diploma_current.txt")
with open(out, "w", encoding="utf-8") as f:
    f.write(full)
print("wrote", out)
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    if p.style.name.startswith("Heading") or re.match(r"^(\d+\.?\d*)\s", t):
        print(p.style.name, t[:90])
