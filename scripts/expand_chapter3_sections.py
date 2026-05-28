# -*- coding: utf-8 -*-
"""Добавляет подпункты главы 3 со слотами для скриншотов (без изображений)."""
from __future__ import annotations

import glob
import os
import shutil
from datetime import datetime

from docx import Document

from chapter3_extra_sections import CHAPTER3_INSERTIONS, INTRO_FIGURES_SUFFIX
from insert_chapter2 import apply_body_format, docx_paths, insert_blocks_before

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def find_paragraph_starting(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"Не найден абзац с текстом: {prefix!r}")


def already_expanded(doc: Document) -> bool:
    return any(p.text.strip().startswith("3.1.1 Навигация") for p in doc.paragraphs)


def patch_chapter3_intro(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("В разделе описано фактическое поведение"):
            if "3.6–3.17" in p.text:
                return
            base = p.text.rstrip()
            for old in (
                " Ниже приведены скриншоты основных экранов (рисунки 3.1–3.5).",
                " Скриншоты основных экранов приведены на рисунках 3.1–3.5;",
            ):
                base = base.replace(old, "")
            if "опирается на исходный код" in base and not base.endswith("."):
                base += "."
            p.text = base.rstrip(".") + "." + INTRO_FIGURES_SUFFIX
            return


def main() -> None:
    paths = docx_paths()
    if not paths:
        raise SystemExit("Закройте Word. Файл Дипломначалоконец.docx не найден.")
    path = [p for p in paths if "_updated" not in p and "_synced" not in p][0]

    doc = Document(path)
    if already_expanded(doc):
        print("Подпункты 3.1.1 и далее уже добавлены.")
        return

    backup = path + ".bak_ch3exp_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    body = next(
        (p for p in doc.paragraphs if p.style.name == "Normal" and len(p.text.strip()) > 80),
        doc.paragraphs[29],
    )

    for anchor_prefix, blocks in CHAPTER3_INSERTIONS:
        anchor = find_paragraph_starting(doc, anchor_prefix)
        insert_blocks_before(anchor, blocks, body)

    patch_chapter3_intro(doc)

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_ch3expanded.docx")
        doc.save(alt)
        print("Файл занят. Сохранено:", alt)


if __name__ == "__main__":
    main()
