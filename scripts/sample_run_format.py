# -*- coding: utf-8 -*-
import glob
import os

from docx import Document
from docx.shared import Pt

from diploma_paths import example_docx_path, workspace_file

ex = example_docx_path()
if not ex:
    raise SystemExit("Пример .docx не найден в diploma/references/")
doc = Document(ex)
out = workspace_file("logs", "_run_format.txt")
needles = ["Несмотря на стремительное", "Рисунок 1.3.1", "Листинг 2.5.1"]
with open(out, "w", encoding="utf-8") as f:
    for p in doc.paragraphs:
        if not any(n in p.text for n in needles):
            continue
        pf = p.paragraph_format
        f.write(f"=== {p.text[:65]} ===\n")
        f.write(f"style={p.style.name} align={pf.alignment} indent={pf.first_line_indent}\n")
        f.write(f"line_rule={pf.line_spacing_rule} line={pf.line_spacing}\n")
        if p.runs:
            r = p.runs[0]
            sz = r.font.size
            f.write(f"run0 font={r.font.name} size={sz} pt={sz.pt if sz else None}\n")
        if "Листинг" in p.text:
            for j, pp in enumerate(doc.paragraphs):
                if pp._p is p._p:
                    for k in range(1, 3):
                        cp = doc.paragraphs[j + k]
                        cr = cp.runs[0] if cp.runs else None
                        f.write(
                            f"  +{k} style={cp.style.name} "
                            f"font={cr.font.name if cr else None}\n"
                        )
                    break
        f.write("\n")
print(out)
