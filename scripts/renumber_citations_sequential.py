# -*- coding: utf-8 -*-
"""
Перенумерация ссылок [1], [2], … по порядку первого упоминания в тексте.
Список литературы — только использованные источники, в том же порядке.
"""
from __future__ import annotations

import glob
import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from diploma_sources import BIBLIOGRAPHY_MONOGR

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


from diploma_paths import diploma_path


def find_doc() -> str:
    return diploma_path()


def find_zones(doc: Document) -> tuple[int, int]:
    intro = 0
    bib = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Введение" and p.style.name.startswith("Heading"):
            intro = i
        if "Список использованных источников" in t:
            bib = i
    return intro, bib


def has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def collect_first_use_order(doc: Document, intro: int, bib: int) -> list[int]:
    """Порядок старых номеров [N] по первому появлению в тексте."""
    order: list[int] = []
    seen: set[int] = set()
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p):
            continue
        for m in re.finditer(r"\[(\d+)\]", p.text):
            n = int(m.group(1))
            if 1 <= n <= len(BIBLIOGRAPHY_MONOGR) and n not in seen:
                seen.add(n)
                order.append(n)
    return order


def replace_refs_in_text(text: str, old_to_new: dict[int, int]) -> str:
    def sub(m: re.Match) -> str:
        old = int(m.group(1))
        if old in old_to_new:
            return f"[{old_to_new[old]}]"
        return m.group(0)

    return re.sub(r"\[(\d+)\]", sub, text)


def remove_duplicate_citations(doc: Document, intro: int, bib: int) -> int:
    """Удаляет повторные [N] в тексте — каждый номер только при первом упоминании."""
    cited: set[int] = set()
    removed = 0

    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p) or "[" not in p.text:
            continue

        def sub(m: re.Match) -> str:
            nonlocal removed
            n = int(m.group(1))
            if n in cited:
                removed += 1
                return ""
            cited.add(n)
            return m.group(0)

        new_t = re.sub(r"\s*\[(\d+)\]", sub, p.text)
        new_t = re.sub(r"  +", " ", new_t)
        new_t = re.sub(r" +([.,;])", r"\1", new_t)
        if new_t != p.text:
            set_paragraph_text_preserve_style(p, new_t)

    return removed


def set_paragraph_text_preserve_style(paragraph: Paragraph, new_text: str) -> None:
    """Меняет текст, сохраняя стиль (через run[0] если один run)."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = new_text
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def apply_bib_style(paragraph: Paragraph) -> None:
    try:
        paragraph.style = "Документация"
    except KeyError:
        pass
    p_pr = paragraph._element.get_or_add_pPr()
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), "a0")
    old = p_pr.find(qn("w:pStyle"))
    if old is not None:
        p_pr.remove(old)
    p_pr.insert(0, p_style)
    pf = paragraph.paragraph_format
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = pf.left_indent = None


def replace_bibliography(doc: Document, bib_start: int, ordered_old_ids: list[int]) -> None:
    to_remove: list = []
    i = bib_start + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            i += 1
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        to_remove.append(p)
        i += 1

    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    heading = doc.paragraphs[bib_start]
    last = heading
    for old_id in ordered_old_ids:
        entry = BIBLIOGRAPHY_MONOGR[old_id - 1]
        new_p = insert_paragraph_after(last, entry)
        apply_bib_style(new_p)
        last = new_p


def main() -> None:
    path = find_doc()
    backup = path + ".bak_seq_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    intro, bib = find_zones(doc)

    first_order = collect_first_use_order(doc, intro, bib)
    if not first_order:
        print("Ссылки [N] в тексте не найдены.")
        return

    old_to_new = {old: i + 1 for i, old in enumerate(first_order)}
    print("Порядок первого упоминания (старый -> новый):")
    for old in first_order:
        print(f"  [{old}] -> [{old_to_new[old]}]")

    changed = 0
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p) or not p.text:
            continue
        new_t = replace_refs_in_text(p.text, old_to_new)
        if new_t != p.text:
            set_paragraph_text_preserve_style(p, new_t)
            changed += 1

    replace_bibliography(doc, bib, first_order)

    removed = remove_duplicate_citations(doc, intro, bib)
    if removed:
        print(f"Удалено повторных ссылок в тексте: {removed}")

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = path.replace(".docx", "_sequential.docx")
        doc.save(alt)
        print("Файл занят. Сохранено:", alt)

    print(f"Обновлено абзацев: {changed}, источников в списке: {len(first_order)}")


if __name__ == "__main__":
    main()
