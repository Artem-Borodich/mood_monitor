# -*- coding: utf-8 -*-
"""Синхронизация Дипломначалоконец.docx с репозиторием mood_project."""
from __future__ import annotations

import glob
import os
import shutil
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph

from insert_chapter2 import (
    apply_body_format,
    docx_paths,
    insert_blocks_before,
    insert_paragraph_before,
)
from sync_diploma_content import (
    CONCLUSION_REPLACE,
    INTRO_REPLACE_LOCAL_STORAGE,
    SECTION_14_BLOCKS,
    SECTION_24_BLOCKS,
    SECTION_25_BLOCKS,
    SECTION_3_BLOCKS,
    SUBJECT_RESEARCH,
)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def collect_between(
    doc: Document, start_prefix: str, end_prefix: str, include_start: bool = False
) -> tuple[list[Paragraph], Paragraph | None]:
    to_remove: list[Paragraph] = []
    anchor: Paragraph | None = None
    active = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(start_prefix):
            active = True
            if include_start:
                to_remove.append(p)
            continue
        if active:
            if end_prefix and t.startswith(end_prefix):
                anchor = p
                break
            to_remove.append(p)
    return to_remove, anchor


def replace_section(
    doc: Document,
    start_prefix: str,
    end_prefix: str,
    blocks: list[tuple[str, str]],
    body_template: Paragraph,
    include_start: bool = True,
) -> None:
    if include_start:
        to_remove, anchor = collect_between(doc, start_prefix, end_prefix, True)
    else:
        to_remove, anchor = collect_between(doc, start_prefix, end_prefix, False)
    if anchor is None:
        raise RuntimeError(f"Якорь не найден: {end_prefix}")
    for p in reversed(to_remove):
        remove_paragraph(p)
    insert_blocks_before(anchor, blocks, body_template)


def patch_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text
        if "локальных баз данных" in t and "устройстве пользователя" in t:
            p.text = INTRO_REPLACE_LOCAL_STORAGE[1]
        elif t.startswith("Локальное хранение данных позволяет пользователю сохранять историю"):
            p.text = (
                "Серверное хранение обеспечивает единый журнал для всех сессий клиента; "
                "накопленные записи используются для графиков, индекса, рекомендаций и прогноза."
            )
        elif t.startswith("Предмет исследования:"):
            p.text = SUBJECT_RESEARCH[1]
        elif (
            "общей картины эмоционального состояния пользователя" in t
            and "за определённый период" in t
        ):
            p.text = (
                "Разрабатываемое приложение — программная система для фиксации "
                "настроения, стресса и энергии с хранением журнала в PostgreSQL и "
                "доступом через программный интерфейс FastAPI. Пользователь добавляет "
                "записи с опциональными полями сна, активности, категории и заметки; "
                "сервер вычисляет индекс благополучия, рекомендации и прогноз, клиент "
                "отображает графики, историю и советы."
            )
        elif "условный индекс благополучия" in t and "за определённый период" in t:
            for old in (
                "за определённый период времени",
                "за определённый период",
            ):
                if old in t:
                    p.text = t.replace(
                        old,
                        "по последней записи в журнале (запрос к серверу)",
                    )
                    break
        elif "Все данные сохраняются локально" in t:
            p.text = CONCLUSION_REPLACE[1]
        elif "Диаграмму вариантов использования" in t and "Рисунок" in t:
            p.text = (
                "Варианты использования согласованы с реализацией и отражены "
                "в §2.3.1 и в требованиях §1.3.4."
            )
        elif "PostgreSQL (не PostgreSQL)" in t:
            p.text = t.replace("(не PostgreSQL)", "(не SQLite)")
def main() -> None:
    paths = docx_paths()
    if not paths:
        raise SystemExit("Закройте Word. Файл Дипломначалоконец.docx не найден.")
    path = [p for p in paths if "_updated" not in p][0]

    backup = path + ".bak_sync_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    body = next(
        (p for p in doc.paragraphs if p.style.name == "Normal" and len(p.text.strip()) > 80),
        doc.paragraphs[29],
    )

    patch_paragraphs(doc)

    replace_section(doc, "1.4 Выбор средств", "1.5 Выбор средств", SECTION_14_BLOCKS, body)

    # 2.4 + вставка 2.5 перед главой 3
    blocks_24_25 = SECTION_24_BLOCKS + SECTION_25_BLOCKS
    replace_section(
        doc,
        "2.4 Программная реализация",
        "3 Интерфейс",
        blocks_24_25,
        body,
    )

    replace_section(doc, "3 Интерфейс", "Заключение", SECTION_3_BLOCKS, body, include_start=True)

    # Заключение: ещё один абзац про локальное хранение
    for p in doc.paragraphs:
        if "сохраняются локально и могут быть просмотрены" in p.text:
            p.text = CONCLUSION_REPLACE[1]

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_synced.docx")
        doc.save(alt)
        print("Файл занят. Сохранено:", alt)


if __name__ == "__main__":
    main()
