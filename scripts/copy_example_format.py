# -*- coding: utf-8 -*-
"""
Оформление по примеру Нестерович: стиль Документация, листинги, ссылки,
библиография monogr., колонтитулы. Титул и реферат не трогаем.
"""
from __future__ import annotations

import glob
import os
import re
import shutil
import zipfile
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from diploma_sources import BIBLIOGRAPHY_MONOGR

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STYLE_BODY = "Документация"
STYLE_CODE = "ConsPlusNonformat"
INTRO_MARKER = "Введение"


from diploma_paths import diploma_path, example_docx_path


def paths() -> tuple[str, str | None]:
    mine = diploma_path()
    return mine, example_docx_path()


def has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def intro_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == INTRO_MARKER and p.style.name.startswith("Heading"):
            return i
    raise RuntimeError("Введение не найдено")


def find_zones(doc: Document) -> tuple[int, int, int]:
    intro = intro_index(doc)
    bib = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        if "Список использованных источников" in p.text:
            bib = i
    return intro, bib, len(doc.paragraphs)


def copy_style_from_example(example_path: str, target_path: str) -> None:
    from copy_style_xml import merge_styles

    merge_styles(example_path, target_path, ["a0", "ConsPlusNonformat"])


def ensure_styles_in_document(doc: Document) -> None:
    """Проверка наличия стилей после копирования XML."""
    names = {s.name for s in doc.styles}
    if STYLE_BODY not in names:
        # fallback: переименовать через прямое обращение к latent styles не получится
        raise RuntimeError(
            f"Стиль {STYLE_BODY!r} недоступен. Откройте документ в Word и сохраните."
        )


def set_paragraph_style(paragraph: Paragraph, style_name: str, style_id: str | None = None) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        sid = style_id or ("a0" if style_name == STYLE_BODY else "ConsPlusNonformat")
        p_pr = paragraph._element.get_or_add_pPr()
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), sid)
        old = p_pr.find(qn("w:pStyle"))
        if old is not None:
            p_pr.remove(old)
        p_pr.insert(0, p_style)


def is_caption(text: str) -> bool:
    t = text.strip()
    return t.startswith("Рисунок") or t.startswith("Таблица")


def is_listing_title(text: str) -> bool:
    return bool(re.match(r"^Листинг\s+\d+\.\d+", text.strip()))


def is_code_paragraph(paragraph: Paragraph, text: str) -> bool:
    if paragraph.style.name == "Normal (Web)":
        return True
    markers = ("@router", "def ", "import ", "class ", "return ", "engine =", "SessionLocal")
    return any(m in text for m in markers)


def fix_listing_titles(doc: Document, intro: int, bib: int) -> int:
    """Листинг 2.5 — ... → Листинг 2.5.1 – ..."""
    count = 0
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        t = p.text.strip()
        m = re.match(r"^Листинг\s+(\d+)\.(\d+)\s*[—\-–]\s*(.+)$", t)
        if not m:
            m = re.match(r"^Листинг\s+(\d+)\.(\d+)\.\d+\s*[—\-–]\s*(.+)$", t)
            if m:
                set_paragraph_style(p, STYLE_BODY, "a0")
                continue
        if m:
            ch, num, rest = m.group(1), m.group(2), m.group(3)
            new_t = f"Листинг {ch}.{num}.1 – {rest}"
            if new_t != t:
                p.text = new_t
                set_paragraph_style(p, STYLE_BODY)
                count += 1
        m2 = re.match(r"^Листинг\s+(\d+)\.(\d+)\.\d+\s*[—\-]\s*", t)
        if m2:
            set_paragraph_style(p, STYLE_BODY)
    return count


def strip_all_refs(text: str) -> str:
    return re.sub(r"\s*\[\d+\]", "", text)


def curated_refs() -> list[tuple[list[str], int]]:
    """Точечные ссылки: не более одного срабатывания на абзац по первому ключу."""
    return [
        (["Flutter"], 1),
        (["Dart"], 2),
        (["Python"], 3),
        (["FastAPI"], 4),
        (["SQLAlchemy"], 5),
        (["docs.flutter.dev"], 6),
        (["fastapi.tiangolo.com"], 7),
        (["Material Design"], 8),
        (["PostgreSQL"], 9),
        (["архитектурные стили", "RESTful"], 10),
        (["REST API", "RESTful Web"], 11),
        (["Компьютерные сети", "сетев"], 12),
        (["мобильная разработка", "мобильн"], 12),
        (["клиент-сервер", "клиентской и серверной"], 14),
        (["психодиагностик"], 15),
        (["мониторинг эмоционального", "Программные продукты"], 16),
        (["docs.sqlalchemy.org"], 17),
        (["Pydantic", "pydantic"], 18),
        (["объектно-ориентированного"], 19),
        (["Чистая архитектура", "чистая архитектура"], 20),
    ]


def add_single_ref(paragraph: Paragraph) -> int | None:
    if has_image(paragraph):
        return None
    text = strip_all_refs(paragraph.text)
    if not text.strip() or paragraph.style.name.startswith("Heading"):
        return None
    if is_caption(text) or is_listing_title(text):
        return None
    lower = text.lower()
    for keywords, num in curated_refs():
        if any(kw.lower() in lower for kw in keywords):
            if f"[{num}]" in paragraph.text:
                return None
            stripped = text.rstrip()
            if stripped.endswith("."):
                paragraph.text = stripped[:-1] + f" [{num}]."
            else:
                paragraph.text = stripped + f" [{num}]"
            return num
    return None


def apply_body_style(paragraph: Paragraph, centered: bool = False) -> None:
    set_paragraph_style(paragraph, STYLE_BODY, "a0")
    pf = paragraph.paragraph_format
    if centered:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5


def apply_code_style(paragraph: Paragraph) -> None:
    set_paragraph_style(paragraph, STYLE_CODE, "ConsPlusNonformat")
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def format_paragraphs(doc: Document, intro: int, bib: int) -> None:
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            continue
        if p.style.name.startswith("Heading"):
            continue
        if is_code_paragraph(p, t):
            apply_code_style(p)
            continue
        if is_caption(t):
            apply_body_style(p, centered=True)
            continue
        if is_listing_title(t) or t.startswith("Листинг"):
            apply_body_style(p, centered=False)
            continue
        apply_body_style(p, centered=False)


def clean_and_add_references(doc: Document, intro: int, bib: int) -> dict:
    used: set[int] = set()
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p) or not p.text.strip():
            continue
        if p.style.name.startswith("Heading"):
            continue
        cleaned = strip_all_refs(p.text)
        if cleaned != p.text:
            p.text = cleaned

    for i in range(intro, bib):
        num = add_single_ref(doc.paragraphs[i])
        if num:
            used.add(num)

    # гарантия: вставить по одной ссылке в ключевые абзацы для неиспользованных
    anchors = {
        6: "компонентов интерфейса",
        7: "FastAPI",
        10: "Филдинг",
        11: "HTTP-запрос",
        13: "кроссплатформен",
        14: "клиент-сервер",
        15: "психодиагностик",
        16: "мониторинг",
        17: "SQLAlchemy",
        19: "проектирован",
        20: "архитектур",
    }
    for num, kw in anchors.items():
        if num in used:
            continue
        for i in range(intro, bib):
            p = doc.paragraphs[i]
            if kw.lower() in p.text.lower() and "[" not in p.text:
                t = p.text.rstrip()
                p.text = (t[:-1] if t.endswith(".") else t) + f" [{num}]."
                used.add(num)
                break

    for num, kw in {
        6: "флаттер",
        7: "fastapi",
        13: "android",
        15: "самочувств",
        17: "orm",
        19: "паттерн",
    }.items():
        if num in used:
            continue
        for i in range(intro, bib):
            p = doc.paragraphs[i]
            if kw in p.text.lower() and f"[{num}]" not in p.text:
                t = strip_all_refs(p.text).rstrip()
                p.text = (t[:-1] if t.endswith(".") else t) + f" [{num}]."
                used.add(num)
                break

    return {"used": sorted(used), "count": len(used)}


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def replace_bibliography(doc: Document, bib_start: int) -> None:
    heading = doc.paragraphs[bib_start]
    set_paragraph_style(heading, "Heading 1")

    to_remove: list = []
    i = bib_start + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            i += 1
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        to_remove.append(p)
        i += 1

    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    last = heading
    for entry in BIBLIOGRAPHY_MONOGR:
        new_p = insert_paragraph_after(last, entry)
        apply_body_style(new_p, centered=False)
        new_p.paragraph_format.first_line_indent = Cm(0)
        new_p.paragraph_format.left_indent = Cm(0)
        last = new_p


def move_bibliography_after_conclusion(doc: Document) -> bool:
    """Если список литературы перед заключением — перенести после него."""
    concl_idx = bib_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Заключение" and p.style.name.startswith("Heading"):
            concl_idx = i
        if "Список использованных источников" in t:
            bib_idx = i
    if concl_idx is None or bib_idx is None or bib_idx > concl_idx:
        return False

    # собрать блок библиографии
    elems = [doc.paragraphs[bib_idx]._element]
    i = bib_idx + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.style.name.startswith("Heading") and "Список" not in p.text:
            break
        if p.text.strip().startswith("Заключение"):
            break
        elems.append(p._element)
        i += 1

    for el in elems:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # вставить после заключения (перед следующим heading)
    insert_after = doc.paragraphs[concl_idx]._element
    j = concl_idx + 1
    while j < len(doc.paragraphs):
        p = doc.paragraphs[j]
        if p.style.name.startswith("Heading") and p.text.strip() != "":
            insert_after = doc.paragraphs[j - 1]._element
            break
        insert_after = p._element
        j += 1

    for el in reversed(elems):
        insert_after.addnext(el)
    return True


def main() -> None:
    mine, ex = paths()
    backup = mine + ".bak_format2_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(mine, backup)
    print("Резервная копия:", backup)

    copy_style_from_example(ex, mine)

    doc = Document(mine)
    try:
        ensure_styles_in_document(doc)
    except RuntimeError as e:
        print("Предупреждение:", e)

    intro, bib, _ = find_zones(doc)
    print(f"Зона правки: абзацы {intro}..{bib - 1} (титул 0..{intro - 1} не трогаем)")

    moved = move_bibliography_after_conclusion(doc)
    if moved:
        intro, bib, _ = find_zones(doc)
        print("Список литературы перенесён после заключения")

    listings = fix_listing_titles(doc, intro, bib)
    format_paragraphs(doc, intro, bib)
    ref_stats = clean_and_add_references(doc, intro, bib)

    intro, bib, _ = find_zones(doc)
    replace_bibliography(doc, bib)

    try:
        doc.save(mine)
        print("Сохранено:", mine)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_formatted.docx")
        doc.save(alt)
        mine = alt
        print("Файл занят. Сохранено:", alt)

    print("Листингов исправлено:", listings)
    try:
        from merge_footers import merge_footers

        merge_footers(ex, mine)
        print("Колонтитулы скопированы из примера")
    except Exception as exc:
        print("Колонтитулы (при ошибке: python scripts/merge_footers.py):", exc)
    print("Ссылки:", ref_stats)


if __name__ == "__main__":
    main()
