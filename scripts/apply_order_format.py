# -*- coding: utf-8 -*-
"""
Приведение диплома к оформлению по «Порядку оформления работы»
(раздел II, ВГУ им. П.М. Машерова). Титульный лист не изменяется.
"""
from __future__ import annotations

import re
import shutil
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from diploma_paths import diploma_path
from diploma_sources import BIBLIOGRAPHY_MONOGR

FONT = "Times New Roman"
SIZE_BODY = Pt(14)
SIZE_CODE = Pt(10)
LINE_EXACT = Pt(18)
INDENT = Cm(1.25)
INTRO_MARKER = "Введение"
TOC_MARKER = "Оглавление"
STYLE_BODY = "Документация"
STYLE_BODY_ID = "a0"


def has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def intro_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == INTRO_MARKER and p.style.name.startswith("Heading"):
            return i
    raise RuntimeError("Введение не найдено")


def find_bib_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "Список использованных источников" in p.text:
            return i
    return len(doc.paragraphs)


def has_toc(doc: Document, before: int) -> bool:
    for i in range(before):
        t = doc.paragraphs[i].text.strip()
        if t in (TOC_MARKER, "Содержание"):
            return True
        if "TOC" in doc.paragraphs[i]._element.xml and "instrText" in doc.paragraphs[i]._element.xml:
            return True
    return False


def set_run_font(run, size=SIZE_BODY, bold=False, italic=False, name=FONT) -> None:
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)


def set_paragraph_style(paragraph: Paragraph, style_name: str, style_id: str | None = None) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        sid = style_id or STYLE_BODY_ID
        p_pr = paragraph._element.get_or_add_pPr()
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), sid)
        old = p_pr.find(qn("w:pStyle"))
        if old is not None:
            p_pr.remove(old)
        p_pr.insert(0, p_style)


def apply_body_format(paragraph: Paragraph) -> None:
    set_paragraph_style(paragraph, STYLE_BODY, STYLE_BODY_ID)
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = INDENT
    pf.left_indent = pf.right_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_EXACT
    pf.space_before = pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def apply_heading_format(paragraph: Paragraph, level: int) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = INDENT
    pf.left_indent = pf.right_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_EXACT
    pf.space_before = LINE_EXACT
    pf.space_after = Pt(0)
    text = paragraph.text.strip().lstrip("\t").rstrip(".")
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def apply_figure_caption(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    text = re.sub(r"^(Рисунок\s+[\d.]+)\s*[-–—]\s*", r"\1 – ", text)
    if text != paragraph.text.strip():
        if paragraph.runs:
            paragraph.runs[0].text = text
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.text = text
    set_paragraph_style(paragraph, STYLE_BODY, STYLE_BODY_ID)
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = pf.left_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_EXACT
    pf.space_before = pf.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def apply_table_caption(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    if not re.match(r"^Таблица\s+\d", text, re.I):
        # Нумерация в пределах раздела — только если номера ещё нет
        m = re.match(r"^Таблица\s+(.+)$", text, re.I)
        if m:
            text = f"Таблица 2.1 – {m.group(1)}"
    text = re.sub(r"^(Таблица\s+[\d.]+)\s*[-–—]?\s*", r"\1 – ", text)
    if text != paragraph.text.strip():
        if paragraph.runs:
            paragraph.runs[0].text = text
            for r in paragraph.runs[1:]:
                r.text = ""
    set_paragraph_style(paragraph, STYLE_BODY, STYLE_BODY_ID)
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = pf.left_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_EXACT
    pf.space_before = pf.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def apply_listing_caption(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    text = re.sub(r"^(Листинг\s+[\d.]+)\s*[-–—]\s*", r"\1 – ", text)
    if text != paragraph.text.strip() and paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    apply_body_format(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0)


def apply_code_format(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = pf.left_indent = Cm(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=SIZE_CODE, name="Courier New")


def apply_bibliography_entry(paragraph: Paragraph, number: int, text: str) -> None:
    paragraph.text = f"{number} {text}"
    apply_body_format(paragraph)
    pf = paragraph.paragraph_format
    pf.first_line_indent = INDENT
    pf.left_indent = None


def heading_level(style_name: str) -> int | None:
    if style_name == "Heading 1":
        return 1
    if style_name == "Heading 2":
        return 2
    if style_name == "Heading 3":
        return 3
    return None


def is_caption(text: str) -> bool:
    t = text.strip()
    return t.startswith("Рисунок")


def is_table_caption(text: str) -> bool:
    return text.strip().startswith("Таблица")


def is_listing_caption(text: str) -> bool:
    return text.strip().startswith("Листинг")


def is_code_paragraph(paragraph: Paragraph, text: str) -> bool:
    if paragraph.style.name == "Normal (Web)":
        return True
    markers = ("@router", "def ", "import ", "class ", "return ", "engine =", "SessionLocal", "CREATE TABLE")
    return any(m in text for m in markers)


def insert_paragraph_before(ref: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    ref._p.addprevious(new_p)
    para = Paragraph(new_p, ref._parent)
    if text:
        para.add_run(text)
    return para


def add_page_break_paragraph(ref: Paragraph) -> Paragraph:
    p = insert_paragraph_before(ref)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_toc_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def insert_toc_before_intro(doc: Document, intro: int) -> bool:
    if has_toc(doc, intro):
        return False
    intro_p = doc.paragraphs[intro]
    add_page_break_paragraph(intro_p)
    toc_field = insert_paragraph_before(intro_p)
    add_toc_field(toc_field)
    toc_heading = insert_paragraph_before(toc_field, TOC_MARKER)
    toc_heading.style = "Heading 1"
    apply_heading_format(toc_heading, 1)
    add_page_break_paragraph(toc_heading)
    return True


def update_document_styles(doc: Document) -> None:
    for name in ("Normal", STYLE_BODY):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.size = SIZE_BODY
        pf = st.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = INDENT
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = LINE_EXACT

    for level in (1, 2, 3):
        try:
            st = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.size = SIZE_BODY
        st.font.bold = True
        pf = st.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = INDENT
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = LINE_EXACT
        pf.space_before = LINE_EXACT
        pf.space_after = Pt(0)


def set_section_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.different_first_page_header_footer = True
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), "1")


def _add_page_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)


def setup_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
        fp.text = ""

        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        set_run_font(run)
        _add_page_field(run)


def collect_first_use_order(doc: Document, intro: int, bib: int) -> list[int]:
    order: list[int] = []
    seen: set[int] = set()
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p):
            continue
        for m in re.finditer(r"\[(\d+)\]", p.text):
            n = int(m.group(1))
            if 1 <= n <= len(BIBLIOGRAPHY_MONOGR) and n not in seen:
                seen.add(n)
                order.append(n)
    return order


def replace_refs_in_text(text: str, old_to_new: dict[int, int]) -> str:
    def sub(m: re.Match) -> str:
        old = int(m.group(1))
        return f"[{old_to_new[old]}]" if old in old_to_new else m.group(0)

    return re.sub(r"\[(\d+)\]", sub, text)


def set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def rebuild_bibliography(doc: Document, bib_start: int, ordered_old_ids: list[int]) -> None:
    to_remove: list = []
    i = bib_start + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            i += 1
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        if t.startswith("Приложение"):
            break
        to_remove.append(p)
        i += 1

    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    heading = doc.paragraphs[bib_start]
    apply_heading_format(heading, 1)
    last = heading
    for new_num, old_id in enumerate(ordered_old_ids, start=1):
        entry = BIBLIOGRAPHY_MONOGR[old_id - 1]
        new_p = insert_paragraph_after(last)
        apply_bibliography_entry(new_p, new_num, entry)
        last = new_p


def format_content(doc: Document, start: int, end: int) -> dict[str, int]:
    stats = {"body": 0, "headings": 0, "figures": 0, "tables": 0, "listings": 0, "code": 0}
    for i in range(start, end):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t and not has_image(p):
            continue
        lvl = heading_level(p.style.name)
        if lvl:
            apply_heading_format(p, lvl)
            stats["headings"] += 1
            continue
        if is_caption(t):
            apply_figure_caption(p)
            stats["figures"] += 1
            continue
        if is_table_caption(t):
            apply_table_caption(p)
            stats["tables"] += 1
            continue
        if is_listing_caption(t):
            apply_listing_caption(p)
            stats["listings"] += 1
            continue
        if is_code_paragraph(p, t):
            apply_code_format(p)
            stats["code"] += 1
            continue
        if has_image(p):
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            continue
        apply_body_format(p)
        stats["body"] += 1
    return stats


def renumber_citations(doc: Document, intro: int, bib: int) -> tuple[dict[int, int], list[int]]:
    order = collect_first_use_order(doc, intro, bib)
    if not order:
        return {}, []
    old_to_new = {old: i + 1 for i, old in enumerate(order)}
    for i in range(intro, bib):
        p = doc.paragraphs[i]
        if has_image(p) or "[" not in p.text:
            continue
        new_t = replace_refs_in_text(p.text, old_to_new)
        if new_t != p.text:
            set_paragraph_text(p, new_t)
    return old_to_new, order


def main() -> None:
    path = diploma_path()
    backup = path + ".bak_order_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    intro = intro_index(doc)
    print(f"Титульный лист: абзацы 0..{intro - 1} (не изменяются)")

    toc_added = insert_toc_before_intro(doc, intro)
    if toc_added:
        intro = intro_index(doc)
        print("Добавлено оглавление с полем TOC перед «Введение»")

    update_document_styles(doc)
    set_section_layout(doc)
    setup_page_numbers(doc)

    bib = find_bib_index(doc)
    _, ordered = renumber_citations(doc, intro, bib)
    stats = format_content(doc, intro, len(doc.paragraphs))

    bib = find_bib_index(doc)
    if ordered:
        rebuild_bibliography(doc, bib, ordered)
        print(f"Список литературы: {len(ordered)} источников по порядку ссылок")
    else:
        apply_heading_format(doc.paragraphs[bib], 1)

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = path.replace(".docx", "_order_formatted.docx")
        doc.save(alt)
        print("Файл занят Word. Сохранено:", alt)

    print("Статистика:", stats)
    if toc_added:
        print("В Word нажмите F9 для обновления оглавления.")


if __name__ == "__main__":
    main()
