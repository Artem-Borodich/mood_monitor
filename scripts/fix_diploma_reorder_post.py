# -*- coding: utf-8 -*-
"""Доправки: §2.1 перед текстом; §1.3–1.4 как Heading 2."""
from __future__ import annotations

import shutil
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph

from diploma_paths import diploma_path

CHAR_H2 = "2.1 Характеристика разрабатываемого приложения"
CH2_PREFIX = "2 Проектирование и реализация"


def find_para(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def move_element_before(doc: Document, move_p: Paragraph, before_p: Paragraph) -> None:
    el = move_p._p
    body = doc.element.body
    body.remove(el)
    before_p._p.addprevious(el)


def main() -> None:
    path = diploma_path()
    backup = path + ".bak_fix_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)

    doc = Document(path)
    ch2 = None
    char_h = find_para(doc, CHAR_H2)
    for p in doc.paragraphs:
        if p.text.strip().startswith(CH2_PREFIX):
            ch2 = p
            break

    if ch2 and char_h:
        first_body: Paragraph | None = None
        after_ch2 = False
        for p in doc.paragraphs:
            if p._p is ch2._p:
                after_ch2 = True
                continue
            if not after_ch2:
                continue
            if p._p is char_h._p:
                break
            t = p.text.strip()
            if t and "Heading" not in p.style.name:
                first_body = p
                break
        if first_body:
            move_element_before(doc, char_h, first_body)

    for new in (
        "1.3 Выбор средств для разработки клиентской части приложения",
        "1.4 Выбор средств для разработки серверной части приложения",
    ):
        p = find_para(doc, new)
        if not p:
            p = find_para(doc, new.split(" ", 1)[1])
        if p:
            p.text = new
            try:
                p.style = "Heading 2"
            except KeyError:
                pass

    doc.save(path)
    print("Исправлено:", path)


if __name__ == "__main__":
    main()
