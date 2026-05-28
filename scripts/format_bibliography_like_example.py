# -*- coding: utf-8 -*-
"""Оформление списка литературы как в примере (стиль Документация, без нумерации)."""
from __future__ import annotations

import glob
import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STYLE_BODY = "Документация"
STYLE_ID = "a0"
FIRST_INDENT = Cm(1.25)


from diploma_paths import diploma_path


def find_doc() -> str:
    return diploma_path()


def find_bib_start(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "Список использованных источников" in p.text:
            return i
    raise RuntimeError("Раздел списка литературы не найден")


def collect_bib_entries(doc: Document, bib_start: int) -> list[str]:
    entries: list[str] = []
    for i in range(bib_start + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        entries.append(t)
    return entries


def set_style_dokumentatsiya(paragraph: Paragraph) -> None:
    try:
        paragraph.style = STYLE_BODY
    except KeyError:
        pass
    p_pr = paragraph._element.get_or_add_pPr()
    old = p_pr.find(qn("w:pStyle"))
    if old is not None:
        p_pr.remove(old)
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), STYLE_ID)
    p_pr.insert(0, p_style)

    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = FIRST_INDENT
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.italic = False
        run.font.bold = False


def format_heading(paragraph: Paragraph) -> None:
    paragraph.style = "Heading 1"
    pf = paragraph.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True


def normalize_entry(text: str) -> str:
    """Мелкая правка пунктуации под пример."""
    t = text.strip()
    t = re.sub(r"\s+", " ", t)
    t = t.replace(" : [site]. – URL:", " : [site]. – URL:")
    t = re.sub(r"https://www\.espressif\.com /en/", "https://www.espressif.com/en/", t)
    if not t.endswith(".") and "[site]" not in t:
        t += "."
    return t


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def replace_bibliography_section(doc: Document, bib_start: int, entries: list[str]) -> None:
    to_remove: list = []
    i = bib_start + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if not t:
            i += 1
            continue
        if p.style.name.startswith("Heading") and "Список" not in t:
            break
        to_remove.append(p)
        i += 1

    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    heading = doc.paragraphs[bib_start]
    if "Список использованных источников" not in heading.text:
        heading.text = "Список использованных источников"
    format_heading(heading)

    last = heading
    for raw in entries:
        text = normalize_entry(raw)
        new_p = insert_paragraph_after(last, text)
        set_style_dokumentatsiya(new_p)
        last = new_p


def ensure_dokumentatsiya_style(path: str) -> None:
    from copy_style_xml import merge_styles

    from diploma_paths import DIPLOMA_FILENAME

    ex = [
        p
        for p in glob.glob(os.path.join(ROOT, "*.docx"))
        if os.path.basename(p) != DIPLOMA_FILENAME
        and ".bak" not in p
        and "updated" not in p.lower()
        and not os.path.basename(p).startswith("~$")
    ]
    if not ex:
        raise FileNotFoundError("Пример не найден (Нестерович Диплом.docx)")
    ex = ex[0]
    merge_styles(ex, path, ["a0"])


def main() -> None:
    path = find_doc()
    backup = path + ".bak_bibfmt_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    ensure_dokumentatsiya_style(path)

    doc = Document(path)
    bib_start = find_bib_start(doc)
    entries = collect_bib_entries(doc, bib_start)
    if not entries:
        raise RuntimeError("Записи списка литературы не найдены")

    print(f"Записей: {len(entries)}")
    replace_bibliography_section(doc, bib_start, entries)

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = path.replace(".docx", "_bib_fmt.docx")
        doc.save(alt)
        print("Файл занят. Сохранено:", alt)


if __name__ == "__main__":
    main()
