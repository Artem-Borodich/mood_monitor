# -*- coding: utf-8 -*-
import os
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
from diploma_paths import diploma_path

doc = Document(diploma_path())
with open(os.path.join(ROOT, "_structure.txt"), "w", encoding="utf-8") as f:
    for i, p in enumerate(doc.paragraphs[:80]):
        t = p.text.strip()
        if t:
            f.write(f"{i:3} [{p.style.name:18}] {t[:75]}\n")
    f.write("\n...\n")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "Список" in t or t == "Заключение" or t.startswith("Оглавление"):
            f.write(f"{i:3} [{p.style.name:18}] {t[:75]}\n")
