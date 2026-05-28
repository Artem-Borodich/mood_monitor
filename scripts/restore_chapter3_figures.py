# -*- coding: utf-8 -*-
"""Восстановление рисунков 3.1–3.5 в главе 3 из резервной копии диплома."""
from __future__ import annotations

import glob
import os
import shutil
import zipfile
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph

from insert_chapter2 import apply_body_format, docx_paths, insert_blocks_before

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCREENSHOTS_DIR = os.path.join(ROOT, "docs", "diploma_screenshots")
# Резерв с рисунками удалён при cleanup; изображения — в docs/diploma_screenshots/
BACKUP_SOURCE = os.path.join(ROOT, "docs", "diploma_screenshots", "_from_backup.docx")

# image1.png … image5.png в резервной копии соответствуют рисункам 3.1 … 3.5
FIGURES = [
    {
        "anchor_prefix": "3.2 История",
        "image": "screen_3_1.png",
        "source_media": "image1.png",
        "lead": (
            "На рисунке 3.1 показан вид главной панели приложения после запуска "
            "и подключения к серверу."
        ),
        "caption": "Рисунок 3.1 – Вид главной страницы приложения",
    },
    {
        "anchor_prefix": "3.3 Добавление",
        "image": "screen_3_3.png",
        "source_media": "image3.png",
        "lead": (
            "На рисунке 3.3 представлена страница истории с перечнем записей "
            "и режимами просмотра."
        ),
        "caption": "Рисунок 3.3 – Страница истории оценок",
    },
    {
        "anchor_prefix": "3.4 Советы",
        "image": "screen_3_2.png",
        "source_media": "image2.png",
        "lead": (
            "На рисунке 3.2 приведён экран добавления записи о состоянии "
            "(шкалы настроения, стресса и энергии)."
        ),
        "caption": "Рисунок 3.2 – Страница создания записи о состоянии",
    },
    {
        "anchor_prefix": "3.5 Настройки",
        "blocks_extra": [
            (
                "lead",
                "На рисунке 3.4 показана страница советов с персональной "
                "рекомендацией и каталогом материалов.",
            ),
            ("image", "screen_3_4.png", "image4.png"),
            ("caption", "Рисунок 3.4 – Страница советов"),
            (
                "lead",
                "На рисунке 3.5 отображена работа с избранными и сохранёнными "
                "советами.",
            ),
            ("image", "screen_3_5.png", "image5.png"),
            ("caption", "Рисунок 3.5 – Функция сохранённых советов"),
        ],
    },
]


def _needed_png_names() -> list[str]:
    names: list[str] = []
    for item in FIGURES:
        if "blocks_extra" in item:
            for block in item["blocks_extra"]:
                if block[0] == "image":
                    names.append(block[1])
        else:
            names.append(item["image"])
    return names


def ensure_screenshots_extracted() -> None:
    os.makedirs(SCREENSHOTS_DIR, exist_ok=True)
    if all(
        os.path.isfile(os.path.join(SCREENSHOTS_DIR, name))
        for name in _needed_png_names()
    ):
        return
    src = BACKUP_SOURCE
    if not os.path.isfile(src):
        raise FileNotFoundError(
            f"Нет PNG в {SCREENSHOTS_DIR} и нет резерва: {src}"
        )
    with zipfile.ZipFile(src) as z:
        for item in FIGURES:
            if "blocks_extra" in item:
                for block in item["blocks_extra"]:
                    if block[0] == "image":
                        _, dest_name, media_name = block
                        _extract_one(z, media_name, dest_name)
            else:
                _extract_one(z, item["source_media"], item["image"])
        for block in FIGURES[-1]["blocks_extra"]:
            if block[0] == "image":
                pass  # already done


def _extract_one(z: zipfile.ZipFile, media_name: str, dest_name: str) -> str:
    dest = os.path.join(SCREENSHOTS_DIR, dest_name)
    if os.path.isfile(dest):
        return dest
    arc = f"word/media/{media_name}"
    with z.open(arc) as f:
        data = f.read()
    with open(dest, "wb") as out:
        out.write(data)
    return dest


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"Не найден абзац: {prefix!r}")


def figure_already_present(doc: Document) -> bool:
    return any(p.text.strip().startswith("Рисунок 3.1") for p in doc.paragraphs)


def build_blocks(fig: dict) -> list[tuple[str, str]]:
    if "blocks_extra" in fig:
        blocks: list[tuple[str, str]] = []
        for block in fig["blocks_extra"]:
            if block[0] == "lead":
                blocks.append(("Normal", block[1]))
            elif block[0] == "image":
                _, dest_name, _ = block
                blocks.append(("image", os.path.join(SCREENSHOTS_DIR, dest_name)))
            elif block[0] == "caption":
                blocks.append(("caption", block[1]))
        return blocks
    path = os.path.join(SCREENSHOTS_DIR, fig["image"])
    return [
        ("Normal", fig["lead"]),
        ("image", path),
        ("caption", fig["caption"]),
    ]


def main() -> None:
    paths = docx_paths()
    if not paths:
        raise SystemExit("Закройте Word. Файл Дипломначалоконец.docx не найден.")
    path = [p for p in paths if "_updated" not in p and "_synced" not in p][0]

    ensure_screenshots_extracted()

    backup = path + ".bak_fig3_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)
    if figure_already_present(doc):
        print("Рисунки 3.x уже присутствуют — выход.")
        return

    body = next(
        (p for p in doc.paragraphs if p.style.name == "Normal" and len(p.text.strip()) > 80),
        doc.paragraphs[29],
    )

    for fig in FIGURES:
        anchor = find_paragraph(doc, fig["anchor_prefix"])
        insert_blocks_before(anchor, build_blocks(fig), body)

    # Вводный абзац к главе 3 — ссылка на рисунок 3.1
    for p in doc.paragraphs:
        if p.text.strip().startswith("3 Интерфейс"):
            continue
        if p.text.strip().startswith("В разделе описано фактическое поведение"):
            if "рисунок 3.1" not in p.text.lower():
                p.text = (
                    p.text.rstrip()
                    + " Ниже приведены скриншоты основных экранов (рисунки 3.1–3.5)."
                )
            break

    try:
        doc.save(path)
        print("Сохранено:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_figures.docx")
        doc.save(alt)
        print("Файл занят. Сохранено:", alt)


if __name__ == "__main__":
    main()
