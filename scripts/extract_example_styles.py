# -*- coding: utf-8 -*-
import glob
import os
import re

from docx import Document

from diploma_paths import example_docx_path, workspace_file

ex = example_docx_path()
if not ex:
    raise SystemExit("Пример .docx не найден в diploma/references/")
doc = Document(ex)
out_path = workspace_file("styles", "_ex_styles.txt")

with open(out_path, "w", encoding="utf-8") as out:
    found = False
    for p in doc.paragraphs:
        if p.text.strip() == "Введение" and p.style.name == "Heading 1":
            found = True
            continue
        if found and p.style.name == "Документация" and len(p.text) > 80:
            pf = p.paragraph_format
            r = p.runs[0] if p.runs else None
            out.write("BODY: " + p.text[:60] + "...\n")
            out.write(f"  font={r.font.name if r else None} size={r.font.size if r else None}\n")
            out.write(
                f"  indent={pf.first_line_indent} line_rule={pf.line_spacing_rule} "
                f"line={pf.line_spacing}\n"
            )
            out.write(
                f"  align={pf.alignment} space_before={pf.space_before} "
                f"space_after={pf.space_after}\n\n"
            )
            if out.tell() > 1200:
                break

    for p in doc.paragraphs:
        if p.text.startswith("Рисунок 1."):
            r = p.runs[0]
            out.write("CAP: " + p.text[:50] + "\n")
            out.write(
                f"  style={p.style.name} align={p.alignment} "
                f"italic={r.font.italic} size={r.font.size}\n\n"
            )
            break

    for p in doc.paragraphs:
        if "Листинг" in p.text and "–" in p.text:
            out.write("LISTING TITLE: " + p.text[:70] + " style=" + p.style.name + "\n")
            break

    for i, p in enumerate(doc.paragraphs):
        if "Листинг" in p.text and "–" in p.text:
            for j in range(i + 1, i + 4):
                pp = doc.paragraphs[j]
                r = pp.runs[0] if pp.runs else None
                out.write(
                    f"  code para style={pp.style.name} "
                    f"font={r.font.name if r else None} size={r.font.size if r else None}\n"
                )
            break

    for lvl in [1, 2, 3]:
        st = doc.styles[f"Heading {lvl}"]
        out.write(
            f"H{lvl}: font={st.font.name} size={st.font.size} bold={st.font.bold} "
            f"align={st.paragraph_format.alignment}\n"
        )
        out.write(
            f"     space_b={st.paragraph_format.space_before} "
            f"space_a={st.paragraph_format.space_after}\n"
        )

    cnt = 0
    for p in doc.paragraphs:
        if p.style.name == "ConsPlusNonformat" and p.runs:
            r = p.runs[0]
            out.write("CODE: " + p.text[:40] + "\n")
            out.write(f"  font={r.font.name} size={r.font.size}\n")
            cnt += 1
            if cnt >= 2:
                break

    out.write("\nBIB samples:\n")
    in_bib = False
    for p in doc.paragraphs:
        if p.text.strip() == "Список использованных источников":
            in_bib = True
            continue
        if in_bib and p.style.name == "Документация" and p.text.strip():
            out.write(p.text[:120] + "\n")
            if out.tell() > 3500:
                break

print("written", out_path)
