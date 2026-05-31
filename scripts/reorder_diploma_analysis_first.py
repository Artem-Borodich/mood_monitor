# -*- coding: utf-8 -*-
"""
Глава 1 — анализ и исследование; описание приложения — §2.1 (реализация).
"""
from __future__ import annotations

import re
import shutil
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph

from diploma_paths import diploma_path

CH1_TITLE = "1 Анализ и исследование предметной области"
CH2_TITLE = (
    "2 Проектирование и реализация программной системы "
    "мониторинга настроения (интерфейс «Серенити»)"
)
CH3_TITLE = "3 Интерфейс и возможности приложения"
CHAR_H2 = "2.1 Характеристика разрабатываемого приложения"


def find_heading(doc: Document, exact: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    return None


def find_heading_startswith(doc: Document, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix) and "Heading" in p.style.name:
            return p
    return None


def paragraph_before(doc: Document, anchor: Paragraph, skip: Paragraph) -> Paragraph | None:
    prev: Paragraph | None = None
    for p in doc.paragraphs:
        if p._p is anchor._p:
            return prev
        if p._p is not skip._p:
            prev = p
    return None


def collect_range(body, start_el, end_el) -> list:
    out: list = []
    taking = False
    for child in body:
        if child is start_el:
            taking = True
        if taking:
            out.append(child)
        if child is end_el:
            break
    if not out:
        raise RuntimeError("Пустой блок абзацев")
    return out


def move_block_before(doc: Document, start_p: Paragraph, end_p: Paragraph, before_p: Paragraph) -> None:
    body = doc.element.body
    block = collect_range(body, start_p._p, end_p._p)
    for el in block:
        body.remove(el)
    ref = before_p._p
    for el in reversed(block):
        ref.addprevious(el)


def paragraphs_between(doc: Document, start: Paragraph, end: Paragraph) -> list[Paragraph]:
    out: list[Paragraph] = []
    taking = False
    for p in doc.paragraphs:
        if p._p is start._p:
            taking = True
        if taking:
            out.append(p)
        if p._p is end._p:
            break
    return out


def renumber_headings_in(paragraphs: list[Paragraph], pairs: list[tuple[str, str]], skip_exact: set[str] | None = None) -> None:
    skip_exact = skip_exact or set()
    for p in paragraphs:
        if "Heading" not in p.style.name:
            continue
        t = p.text.strip()
        if t in skip_exact:
            continue
        for old, new in pairs:
            if t.startswith(old):
                p.text = new + t[len(old) :]
                break


def fix_cross_refs(doc: Document) -> None:
    repl = [
        ("В §1.1 определено", "В введении определено"),
        ("по §1.1–1.2", "по §1.1–1.2"),
        ("согласованному с §1.1", "согласованному с §2.1"),
        ("требованиям §1.3", "требованиям §1.2"),
        ("Вывод по §1.3.", "Вывод по §1.2."),
        ("§1.3.4", "§1.2.4"),
        ("§1.3.5", "§1.2.5"),
        ("§1.3.6", "§1.2.6"),
        ("§1.3.7", "§1.2.7"),
        ("§1.3.8", "§1.2.8"),
        ("§1.3.1", "§1.2.1"),
        ("§1.3.2", "§1.2.2"),
        ("§1.3.3", "§1.2.3"),
        ("§2.3.1", "§2.4.1"),
        ("§2.5.", "§2.6."),
    ]
    for p in doc.paragraphs:
        t = p.text
        if not t.strip():
            continue
        for old, new in repl:
            t = t.replace(old, new)
        p.text = t


def main() -> None:
    path = diploma_path()
    backup = path + ".bak_reorder_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Резервная копия:", backup)

    doc = Document(path)

    ch1_h = find_heading_startswith(doc, "Разработка приложения")
    char_h = find_heading(doc, "Характеристика приложения")
    analogs_h = find_heading(doc, "1.2 Обзор аналогов")
    ch2_h = find_heading_startswith(doc, "2 Проектирование")
    ch3_h = find_heading_startswith(doc, "3 Интерфейс")
    arch_h = find_heading_startswith(doc, "2.1 Архитектура")

    if not all([ch1_h, char_h, analogs_h, ch2_h, ch3_h, arch_h]):
        raise SystemExit("Не найдены якорные заголовки. Закройте Word.")

    end_char = paragraph_before(doc, analogs_h, char_h)
    if end_char is None:
        raise SystemExit("Не найден конец блока характеристики")

    move_block_before(doc, char_h, end_char, arch_h)

    ch1_h = find_heading_startswith(doc, "Разработка приложения")
    ch2_h = find_heading_startswith(doc, "2 Проектирование")
    ch3_h = find_heading_startswith(doc, "3 Интерфейс")

    ch1_h.text = CH1_TITLE

    char_moved = find_heading(doc, "Характеристика приложения")
    if char_moved:
        char_moved.text = CHAR_H2
        try:
            char_moved.style = "Heading 2"
        except KeyError:
            pass

    # --- Перенумерация только внутри главы 1 (между ch1_h и ch2_h) ---
    ch1_paras = paragraphs_between(doc, ch1_h, ch2_h)[:-1]
    ch1_pairs: list[tuple[str, str]] = []
    for sub in range(8, 0, -1):
        ch1_pairs.append((f"1.3.{sub}", f"1.2.{sub}"))
    for sub in range(5, 0, -1):
        ch1_pairs.append((f"1.2.{sub}", f"1.1.{sub}"))
    ch1_pairs.extend([
        ("1.5 ", "1.4 "),
        ("1.4 ", "1.3 "),
        ("1.3 ", "1.2 "),
        ("1.2 ", "1.1 "),
    ])
    renumber_headings_in(ch1_paras, ch1_pairs)

    # --- Перенумерация только внутри главы 2 (между ch2_h и ch3_h) ---
    ch2_paras = paragraphs_between(doc, ch2_h, ch3_h)[:-1]
    ch2_pairs: list[tuple[str, str]] = []
    for sub in range(9, 0, -1):
        ch2_pairs.append((f"2.5.{sub}", f"2.6.{sub}"))
    for sub in range(9, 0, -1):
        ch2_pairs.append((f"2.4.{sub}", f"2.5.{sub}"))
    for sub in range(9, 0, -1):
        ch2_pairs.append((f"2.3.{sub}", f"2.4.{sub}"))
    for sub in range(9, 0, -1):
        ch2_pairs.append((f"2.2.{sub}", f"2.3.{sub}"))
    for sub in range(9, 0, -1):
        ch2_pairs.append((f"2.1.{sub}", f"2.2.{sub}"))
    ch2_pairs.extend([
        ("2.5 ", "2.6 "),
        ("2.4 ", "2.5 "),
        ("2.3 ", "2.4 "),
        ("2.2 ", "2.3 "),
        ("2.1 ", "2.2 "),
    ])
    renumber_headings_in(ch2_paras, ch2_pairs, skip_exact={CHAR_H2})

    ch2_h.text = CH2_TITLE
    ch3_h.text = CH3_TITLE

    fix_cross_refs(doc)

    doc.save(path)
    print("Сохранено:", path)


if __name__ == "__main__":
    main()
