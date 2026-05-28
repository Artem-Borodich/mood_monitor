# -*- coding: utf-8 -*-
import glob
import os
import re

from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
from diploma_paths import diploma_path

doc = Document(diploma_path())
with open(os.path.join(ROOT, "_refs_scan.txt"), "w", encoding="utf-8") as f:
    refs = re.findall(r"\[\d+\]", "\n".join(p.text for p in doc.paragraphs))
    f.write(f"total bracket refs: {len(refs)}\n")
    f.write(f"unique: {sorted(set(int(x.strip('[]')) for x in refs))}\n\n")
    in_bib = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Список использованных" in t:
            in_bib = True
            f.write("BIB:\n")
            continue
        if in_bib:
            if t.startswith("Приложение"):
                break
            if t:
                f.write(t + "\n")
