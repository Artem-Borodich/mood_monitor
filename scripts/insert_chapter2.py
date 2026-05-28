# -*- coding: utf-8 -*-
"""Insert §2.1–2.3 + diagrams into diploma docx; renumber API listings."""
from __future__ import annotations

import glob
import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

try:
    from chapter2_blocks_full import CHAPTER2_BLOCKS
except ImportError:
    from chapter2_blocks import CHAPTER2_BLOCKS  # type: ignore

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


from diploma_paths import docx_paths as _diploma_docx_paths


def docx_paths() -> list[str]:
    return _diploma_docx_paths()


def insert_paragraph_before(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except KeyError:
            new_para.style = "Normal"
    return new_para


def apply_body_format(paragraph: Paragraph, template: Paragraph) -> None:
    pf = paragraph.paragraph_format
    tpf = template.paragraph_format
    pf.alignment = tpf.alignment
    pf.first_line_indent = tpf.first_line_indent
    pf.line_spacing = tpf.line_spacing
    pf.line_spacing_rule = tpf.line_spacing_rule
    pf.space_after = tpf.space_after
    pf.space_before = tpf.space_before
    if paragraph.runs and template.runs:
        paragraph.runs[0].font.name = template.runs[0].font.name
        paragraph.runs[0].font.size = template.runs[0].font.size


def insert_image_before(
    paragraph: Paragraph,
    image_path: str,
    width_cm: float = 15.0,
    template: Paragraph | None = None,
) -> Paragraph:
    new_p = insert_paragraph_before(paragraph, style="Normal")
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    if template:
        apply_body_format(new_p, template)
    return new_p


def insert_caption_before(
    paragraph: Paragraph, text: str, template: Paragraph | None = None
) -> Paragraph:
    new_p = insert_paragraph_before(paragraph, text, "Normal")
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if template and new_p.runs:
        if template.runs and template.runs[0].font.size:
            new_p.runs[0].font.size = template.runs[0].font.size
        new_p.runs[0].italic = True
    return new_p


def insert_blocks_before(
    anchor: Paragraph,
    blocks: list[tuple[str, str]],
    body_template: Paragraph,
) -> None:
    current = anchor
    for kind, payload in reversed(blocks):
        if kind == "image":
            if not os.path.isfile(payload):
                raise FileNotFoundError(payload)
            current = insert_image_before(
                current, payload, width_cm=15.0, template=body_template
            )
        elif kind == "caption":
            current = insert_caption_before(current, payload, body_template)
        else:
            style = kind
            current = insert_paragraph_before(current, payload, style)
            if style == "Normal":
                apply_body_format(current, body_template)


def find_chapter2_anchor(doc: Document) -> Paragraph:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("2 ") and "API" in t and "взаимодейств" in t:
            return p
        if t.startswith("2 Программная реализация"):
            return p
    raise RuntimeError("Chapter 2 heading not found")


def find_first_body_after_ch2(doc: Document, heading: Paragraph) -> Paragraph:
    found = False
    for p in doc.paragraphs:
        if p._p is heading._p:
            found = True
            continue
        if found and p.text.strip():
            return p
    raise RuntimeError("Chapter 2 body not found")


def update_chapter2_title(heading: Paragraph) -> None:
    heading.text = (
        "2 Проектирование и реализация программной системы Mood Monitor (Serenity)"
    )


def renumber_listings(doc: Document) -> None:
    """Renumber code listings only (not UML figures 2.1–2.4)."""
    mapping = [
        ("Листинг 2.4", "Листинг 2.8"),
        ("Листинг 2.3", "Листинг 2.7"),
        ("Листинг 2.2", "Листинг 2.6"),
        ("Листинг 2.1", "Листинг 2.5"),
    ]
    for p in doc.paragraphs:
        for old, new in mapping:
            if old in p.text:
                p.text = p.text.replace(old, new)


def fix_sqlite_listing(doc: Document) -> None:
    for p in doc.paragraphs:
        if "sqlite" in p.text.lower() or "SQLite" in p.text:
            p.text = p.text.replace("sqlite:///./mood_monitor.db", "postgresql+psycopg2://…/mood_db")
            p.text = re.sub(
                r"базе данных SQLite",
                "базе данных PostgreSQL",
                p.text,
                flags=re.IGNORECASE,
            )
            p.text = re.sub(
                r"SQLite",
                "PostgreSQL",
                p.text,
            )


def already_inserted(doc: Document) -> bool:
    return any(p.text.strip().startswith("2.1 Архитектура") for p in doc.paragraphs)


def pick_source_path(paths: list[str]) -> str:
    """Prefer pre-ch2 backup, else main docx."""
    backups = sorted(
        [p for p in paths if ".bak_ch2_" in p],
        reverse=True,
    )
    main = [p for p in paths if "_updated" not in p and ".bak_" not in p]
    if backups:
        return backups[0]
    return main[0] if main else paths[0]


def main() -> None:
    paths = docx_paths()
    if not paths:
        raise SystemExit("Close Word and ensure Дипломначалоконец.docx exists in project root.")
    path = pick_source_path(paths)
    print("Source:", path)

    if already_inserted(Document(path)):
        print("§2.1 already present in source — use a .bak_ch2_ backup or remove section.")
        return

    backup = path + ".bak_ch2_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    shutil.copy2(path, backup)
    print("Backup:", backup)

    doc = Document(path)
    body_template = None
    for p in doc.paragraphs:
        if p.style.name == "Normal" and len(p.text.strip()) > 80:
            body_template = p
            break
    if body_template is None:
        body_template = doc.paragraphs[29]

    ch2_heading = find_chapter2_anchor(doc)
    update_chapter2_title(ch2_heading)
    anchor = find_first_body_after_ch2(doc, ch2_heading)

    insert_blocks_before(anchor, CHAPTER2_BLOCKS, body_template)
    renumber_listings(doc)
    fix_sqlite_listing(doc)

    try:
        doc.save(path)
        print("Updated:", path)
    except PermissionError:
        alt = os.path.join(ROOT, "Диплома_updated.docx")
        doc.save(alt)
        print("Original file is locked (close Word). Saved to:", alt)


if __name__ == "__main__":
    main()
